# app/services/files.py
from __future__ import annotations

import io
from typing import Tuple, Dict, Any, List, Set

import fitz  # PyMuPDF
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT


def _normalize_font_name(name: str) -> str:
    """
    Normaliza nombres de fuentes embebidas en PDF/DOCX a algo comparables
    con tu lista GOOD_FONTS (arial, calibri, helvetica, verdana, etc).
    Ejemplos: 'ABCDEE+ArialMT' -> 'arial', 'Arial-BoldMT' -> 'arial'
    """
    if not name:
        return ""
    f = name.lower()

    # Quitar prefijos de subconjunto "ABCDEE+"
    if "+" in f:
        f = f.split("+", 1)[1]

    # Quitar sufijos típicos de Adobe/PS
    for rm in ("-boldmt", "-italicmt", "-mt", "-psmt", "mt", "psmt"):
        f = f.replace(rm, "")

    # Normalizaciones simples
    f = f.replace("bold", "").replace("italic", "").strip()
    # Mapear variantes comunes a su familia base
    f = f.replace("arial", "arial")
    f = f.replace("helveticaneue", "helvetica")
    f = f.replace("helvetica", "helvetica")
    f = f.replace("calibri", "calibri")
    f = f.replace("verdana", "verdana")
    f = f.replace("georgia", "georgia")
    f = f.replace("timesnewroman", "times new roman")
    f = f.replace("times new roman", "times new roman")
    f = f.replace("inter", "inter")
    f = f.replace("source sans pro", "source sans pro")

    # último saneo
    return f.strip()


# -----------------------
# PDF con PyMuPDF (fitz)
# -----------------------
def extract_pdf(data: bytes) -> Tuple[str, Dict[str, Any]]:
    """
    Devuelve (texto, meta) donde meta incluye:
      - pages: int
      - images: int
      - fonts: list[str]  (familias normalizadas)
    """
    doc = fitz.open(stream=data, filetype="pdf")
    pages = doc.page_count

    text_parts: List[str] = []
    images = 0
    fonts: Set[str] = set()

    for page in doc:
        # Texto "plano"
        text_parts.append(page.get_text("text"))

        # Contar imágenes reales de la página
        images += len(page.get_images(full=True))

        # Extraer fuentes a partir de los spans
        d = page.get_text("dict")
        for block in d.get("blocks", []):
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    fname = _normalize_font_name(span.get("font") or "")
                    if fname:
                        fonts.add(fname)

    doc.close()

    text = "\n".join(text_parts).strip()
    meta = {
        "pages": pages,
        "images": images,
        "fonts": sorted(fonts),
    }
    return text, meta


# -----------------------
# DOCX con python-docx
# -----------------------
def extract_docx(data: bytes) -> Tuple[str, Dict[str, Any]]:
    """
    Devuelve (texto, meta) donde meta incluye:
      - tables: int
      - images: int
      - fonts: list[str] (familias normalizadas si se encuentran)
    """
    bio = io.BytesIO(data)
    doc = Document(bio)

    # Texto (párrafos + celdas de tablas)
    parts: List[str] = []
    for p in doc.paragraphs:
        parts.append(p.text)

    # Texto de tablas (opcional pero útil)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)

    # Contar tablas
    tables = len(doc.tables)

    # Contar imágenes reales del paquete DOCX
    # (relaciones de tipo IMAGE)
    images = sum(1 for r in doc.part.rels.values() if r.reltype == RT.IMAGE)

    # Familia tipográfica: best-effort (runs pueden no tenerla seteada)
    fonts: Set[str] = set()
    for p in doc.paragraphs:
        for run in p.runs:
            name = None
            if run.font and run.font.name:
                name = run.font.name
            elif run._element.rPr is not None and run._element.rPr.rFonts is not None:
                # fallback XML
                rfs = run._element.rPr.rFonts
                name = rfs.ascii or rfs.hAnsi or rfs.cs or rfs.eastAsia
            if name:
                fonts.add(_normalize_font_name(str(name)))

    text = "\n".join([t for t in parts if t is not None]).strip()
    meta = {
        "tables": tables,
        "images": images,
        "fonts": sorted([f for f in fonts if f]),
    }
    return text, meta
